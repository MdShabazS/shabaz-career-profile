#!/usr/bin/env python3
"""Build the master resume as an ATS-friendly PDF and DOCX from one content source.

Visual template follows the user's prior resume: centered navy header, ruled
section headings, right-aligned dates/locations, quantified bullets, one page.
Header shows clickable "LinkedIn" / "GitHub" labels with small icons (no raw URLs);
email is shown and clickable. Facts are from PROFILE.md / the repo.

PDF  -> reportlab (single column, selectable text, clickable links)
DOCX -> python-docx (single column, live hyperlinks, right-tab dates)

Run: python3 scripts/build_resume.py
Keep this content in sync with resume/master-resume.md.
"""
import os

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "resume")
ASSETS = os.path.join(OUT, "assets")
BASENAME = "Mohammed_Shabaz_S_Master_Resume"
LI_ICON = os.path.join(ASSETS, "linkedin.png")
GH_ICON = os.path.join(ASSETS, "github.png")

NAVY = "#1F3A5F"
LINK = "#1F5FA8"

NAME = "MOHAMMED SHABAZ S"
POSITION = "Embedded Engineer  |  Automotive Embedded Systems"
EMAIL = "md.shabaz.2005@gmail.com"
PHONE = "+91 79755 12403"
LINKEDIN_URL = "https://www.linkedin.com/in/shabaz17/"
GITHUB_URL = "https://github.com/MdShabazS"
LOCATION = "Ballari, Karnataka"

SUMMARY = ("Final-year Electronics & Communication Engineering student (CGPA 8.38) focused on embedded "
           "and automotive firmware — ARM Cortex-M microcontroller systems, finite-state machines and "
           "non-blocking real-time design in Embedded C/C++ on ESP32 and STM32 — with a working software "
           "and computer-vision foundation in Python. Internship experience at Nokia and iHelp Robotics.")

SKILLS = [
    ("Languages", "C, C++, Embedded C, Python, SQL"),
    ("Embedded & Firmware", "Microcontrollers, ARM Cortex-M4, finite-state machines, non-blocking real-time design, GPIO, I2C, ADC, timers, sensor interfacing, debugging"),
    ("Platforms & Tools", "ESP32, STM32 (Nucleo-L476RG), STM32CubeIDE, Arduino, Android Studio, Git & GitHub"),
    ("Libraries & Foundations", "OpenCV, TensorFlow Lite  ·  Data Structures & Algorithms, Operating Systems, DBMS, Computer Networks"),
]

# (title, date, org, location, [bullets])
EXPERIENCE = [
    ("Student Intern", "16 September 2026 – Present",
     "Nokia Solutions and Networks India", "Bangalore",
     ["Working as a Student Intern at Nokia."]),
    ("AI Research Intern – Deep Learning & Model Development", "23 March 2026 – 23 August 2026",
     "iHelp Robotics Private Limited  ·  Project: MITRA", "Remote",
     ["Improved live camera-stream reliability in the MITRA assistive Android app across 2 video sources (hardware WiFi/RTSP and phone-camera fallback): cut stream search delay, persisted the last working transport, and added refresh/stall recovery.",
      "Engineered a visual-freeze watchdog plus 3 status/guard mechanisms — stream-status and cloud-status panels and a WebSocket-gated send guard — and fixed a live-refresh loop that recurred after 5 minutes.",
      "Hardened offline voice commands (offline-preferred recognition with model recovery and microphone retry backoff) and verified builds across multiple devices with Gradle test/lint/assemble; delivered 6+ app-side improvements while model and backend stayed with a separate team."]),
    ("Student Intern – Skin Disease Classification", "1–30 June 2026",
     "IEEE EMBS Pune Section", "Remote",
     ["Trained and evaluated a skin-disease image classifier in a 3-member team during a 1-month IEEE EMBS internship — model selection, training, evaluation, and Python image-processing and testing; delivered a working demo-level workflow."]),
]

# (title, repo_text, repo_url, stack, [bullets])
PROJECTS = [
    ("Automotive Body Control Module",
     "github.com/MdShabazS/Automotive-Body-Control-Module-ESP32",
     "https://github.com/MdShabazS/Automotive-Body-Control-Module-ESP32",
     "ESP32, Arduino/C++, Finite-State Machine, GPIO, I2C OLED",
     ["Programmed an ESP32 body-control module with a 3-state (OFF/ACC/ON) ignition FSM driving 6+ vehicle functions: turn indicators, synchronized hazard, brake logic, buzzer feedback and an OLED dashboard.",
      "Architected a non-blocking millis()-based scheduler (zero delay() in the main loop) with 30 ms brake debouncing, throttled 10 Hz I2C OLED refresh, a clean GPIO abstraction and edge-triggered serial logging for deterministic real-time behavior."]),
    ("Smart Wellness Desk Assistant",
     "github.com/MdShabazS/Smart-Wellness-Desk-Assistant",
     "https://github.com/MdShabazS/Smart-Wellness-Desk-Assistant",
     "STM32 Nucleo-L476RG (ARM Cortex-M4), STM32 HAL, STM32CubeIDE, Embedded C  ·  College team project",
     ["Interfaced ultrasonic presence and 12-bit ADC temperature sensors, an I2C SSD1306 OLED and timer-driven buzzer alerts on an STM32 Nucleo-L476RG (ARM Cortex-M4) using STM32 HAL, coordinated by a non-blocking finite-state machine; wired, tested and debugged in STM32CubeIDE as part of a college team."]),
    ("VisionPay — Real-Time Indian Currency Detection",
     "github.com/MdShabazS/visionpay", "https://github.com/MdShabazS/visionpay",
     "Python, OpenCV, TensorFlow Lite, MobileNetV2, Text-to-Speech",
     ["Created an offline, real-time tool that recognizes Indian currency from a live webcam and speaks the result aloud for visually impaired users.",
      "Collected ~400 images per denomination across 6 classes (Rs.10–Rs.500) and trained a MobileNetV2 classifier converted to TensorFlow Lite, reporting ~93% validation accuracy.",
      "Strengthened real-time reliability with 4 mechanisms: confidence/margin gating, temporal smoothing, a background class, and an auto-count mode."]),
]

EDUCATION = {
    "title": "B.E. Electronics & Communication Engineering",
    "date": "Expected 2027",
    "org": "Ballari Institute of Technology and Management (BITM), Ballari",
    "right": "Final Year  ·  CGPA 8.38",
    "extra": "Class 12: 80%     ·     Class 10: 86.88%",
}

LEADERSHIP = [
    "Vice-Chair, IEEE CAS Society, IEEE Student Branch BITM (promoted from Treasurer); Treasurer, BITM Robotics Club.",
    "Google Student Ambassador (2025); Selected Participant, IEEE SPACE 2026 — B.Tech Initiative.",
    "Embedded AI Systems workshop with STMicroelectronics (DigiToad, 2025); Participant, Techzone Nationals 2K25 Hardware Hackathon.",
]

CERTS = ("Embedded Systems — Internshala Trainings (2025)  ·  Programming in C and C++ with AI  ·  "
         "SQL for Data Analytics with AI  ·  Google Cloud Generative AI — SmartBridge/SmartInternz (2025)  ·  "
         "Python Programming — EISystems (2024)")


# ---------------------------------------------------------------- PDF (reportlab)
def build_pdf():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_CENTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
    from reportlab.lib.colors import HexColor
    from reportlab.lib import colors

    path = os.path.join(OUT, BASENAME + ".pdf")
    doc = SimpleDocTemplate(path, pagesize=A4,
                            leftMargin=13 * mm, rightMargin=13 * mm,
                            topMargin=8 * mm, bottomMargin=8 * mm,
                            title="Mohammed Shabaz S - Resume", author="Mohammed Shabaz S")
    navy = HexColor(NAVY)
    dark = HexColor("#1a1a1a")
    W = doc.width

    name_st = ParagraphStyle("name", fontName="Helvetica-Bold", fontSize=16, leading=18,
                             textColor=navy, alignment=TA_CENTER, spaceAfter=1)
    pos_st = ParagraphStyle("pos", fontName="Helvetica", fontSize=10, leading=12,
                            textColor=navy, alignment=TA_CENTER, spaceAfter=2)
    contact_st = ParagraphStyle("contact", fontName="Helvetica", fontSize=8.8, leading=11.5,
                                textColor=dark, alignment=TA_CENTER)
    head_st = ParagraphStyle("head", fontName="Helvetica-Bold", fontSize=10, leading=11.5,
                             textColor=navy, spaceBefore=2, spaceAfter=0.5)
    body_st = ParagraphStyle("body", fontName="Helvetica", fontSize=9.0, leading=10.6, textColor=dark)
    title_st = ParagraphStyle("etitle", fontName="Helvetica-Bold", fontSize=9.4, leading=10.8, textColor=dark)
    date_st = ParagraphStyle("date", fontName="Helvetica", fontSize=8.8, leading=10.8, textColor=dark, alignment=TA_RIGHT)
    org_st = ParagraphStyle("org", fontName="Helvetica-Oblique", fontSize=9.0, leading=10.6, textColor=dark)
    loc_st = ParagraphStyle("loc", fontName="Helvetica-Oblique", fontSize=8.8, leading=10.6, textColor=dark, alignment=TA_RIGHT)
    stack_st = ParagraphStyle("stack", fontName="Helvetica-Oblique", fontSize=8.6, leading=10.2, textColor=dark)
    repo_st = ParagraphStyle("repo", fontName="Helvetica", fontSize=8.4, leading=10.8, textColor=HexColor(LINK), alignment=TA_RIGHT)
    bullet_st = ParagraphStyle("bullet", fontName="Helvetica", fontSize=9.0, leading=10.6, textColor=dark,
                               leftIndent=10, bulletIndent=1, spaceBefore=0.3)

    story = []

    def rule():
        story.append(HRFlowable(width="100%", thickness=0.7, color=navy, spaceBefore=0.5, spaceAfter=1.5))

    def section(name):
        story.append(Paragraph(name, head_st))
        rule()

    def two_col(left_para, right_para):
        t = Table([[left_para, right_para]], colWidths=[W * 0.70, W * 0.30])
        t.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0.5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ]))
        story.append(t)

    def bullets(items):
        for b in items:
            story.append(Paragraph(b, bullet_st, bulletText="•"))

    # header
    story.append(Paragraph(NAME, name_st))
    story.append(Paragraph(POSITION, pos_st))
    sep = "&nbsp;&nbsp;|&nbsp;&nbsp;"
    contact = (f'{LOCATION}{sep}{PHONE}{sep}'
               f'<a href="mailto:{EMAIL}" color="{LINK}">{EMAIL}</a>{sep}'
               f'<img src="{LI_ICON}" width="9" height="9" valign="-1"/>&nbsp;'
               f'<a href="{LINKEDIN_URL}" color="{LINK}">LinkedIn</a>{sep}'
               f'<img src="{GH_ICON}" width="9" height="9" valign="-1"/>&nbsp;'
               f'<a href="{GITHUB_URL}" color="{LINK}">GitHub</a>')
    story.append(Paragraph(contact, contact_st))
    story.append(Spacer(1, 2))
    story.append(HRFlowable(width="100%", thickness=0.7, color=navy, spaceBefore=1, spaceAfter=0))

    section("SUMMARY")
    story.append(Paragraph(SUMMARY, body_st))

    section("TECHNICAL SKILLS")
    for label, val in SKILLS:
        story.append(Paragraph(f"<b>{label}:</b> {val}", body_st))

    section("EXPERIENCE")
    for title, date, org, loc, items in EXPERIENCE:
        two_col(Paragraph(title, title_st), Paragraph(date, date_st))
        two_col(Paragraph(org, org_st), Paragraph(loc, loc_st))
        bullets(items)

    section("PROJECTS")
    for title, repo_txt, repo_url, stack, items in PROJECTS:
        right = Paragraph(f'<a href="{repo_url}" color="{LINK}">GitHub</a>' if repo_url else "", repo_st)
        two_col(Paragraph(title, title_st), right)
        story.append(Paragraph(stack, stack_st))
        bullets(items)

    section("EDUCATION")
    two_col(Paragraph(EDUCATION["title"], title_st), Paragraph(EDUCATION["date"], date_st))
    two_col(Paragraph(EDUCATION["org"], org_st), Paragraph(EDUCATION["right"], loc_st))
    story.append(Paragraph(EDUCATION["extra"], body_st))

    section("LEADERSHIP & ACTIVITIES")
    bullets(LEADERSHIP)

    section("CERTIFICATIONS")
    story.append(Paragraph(CERTS, body_st))

    doc.build(story)
    return path


# ---------------------------------------------------------------- DOCX (python-docx)
def _add_hyperlink(paragraph, url, text, size=9):
    from docx.oxml.shared import OxmlElement, qn
    part = paragraph.part
    r_id = part.relate_to(url,
                          "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                          is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "1F5FA8"); rpr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2))); rpr.append(sz)
    run.append(rpr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def build_docx():
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    NAVY_RGB = RGBColor(0x1F, 0x3A, 0x5F)
    path = os.path.join(OUT, BASENAME + ".docx")
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10)
    sec = doc.sections[0]
    sec.page_height = Inches(11.69); sec.page_width = Inches(8.27)  # A4
    sec.top_margin = sec.bottom_margin = Inches(0.4)
    sec.left_margin = sec.right_margin = Inches(0.5)
    RIGHT_TAB = Inches(7.27)

    def spacing(p, before=0, after=1.5, line=1.0):
        pf = p.paragraph_format
        pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = line

    def center(p):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def heading(text):
        p = doc.add_paragraph(); r = p.add_run(text); r.bold = True; r.font.size = Pt(10.5)
        r.font.color.rgb = NAVY_RGB
        spacing(p, before=6, after=1)
        pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        for k, v in (("w:val", "single"), ("w:sz", "6"), ("w:space", "1"), ("w:color", "1F3A5F")):
            bottom.set(qn(k), v)
        pbdr.append(bottom); pPr.append(pbdr)

    def tabbed(p):
        p.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)

    def bullet(text):
        p = doc.add_paragraph(style="List Bullet"); p.add_run(text); spacing(p, after=1.5)

    # header
    p = doc.add_paragraph(); center(p); r = p.add_run(NAME); r.bold = True; r.font.size = Pt(17)
    r.font.color.rgb = NAVY_RGB; spacing(p, after=0)
    p = doc.add_paragraph(); center(p); r = p.add_run(POSITION); r.font.size = Pt(10.5)
    r.font.color.rgb = NAVY_RGB; spacing(p, after=1)
    p = doc.add_paragraph(); center(p); spacing(p, after=2)
    p.add_run(f"{LOCATION}  |  {PHONE}  |  ").font.size = Pt(9)
    _add_hyperlink(p, f"mailto:{EMAIL}", EMAIL)
    p.add_run("  |  ").font.size = Pt(9)
    p.add_run().add_picture(LI_ICON, height=Pt(9))
    p.add_run(" ").font.size = Pt(9)
    _add_hyperlink(p, LINKEDIN_URL, "LinkedIn")
    p.add_run("  |  ").font.size = Pt(9)
    p.add_run().add_picture(GH_ICON, height=Pt(9))
    p.add_run(" ").font.size = Pt(9)
    _add_hyperlink(p, GITHUB_URL, "GitHub")

    heading("SUMMARY")
    p = doc.add_paragraph(); p.add_run(SUMMARY); spacing(p, after=2)

    heading("TECHNICAL SKILLS")
    for label, val in SKILLS:
        p = doc.add_paragraph(); p.add_run(f"{label}: ").bold = True; p.add_run(val); spacing(p, after=1)

    heading("EXPERIENCE")
    for title, date, org, loc, items in EXPERIENCE:
        p = doc.add_paragraph(); tabbed(p); r = p.add_run(title); r.bold = True; p.add_run("\t" + date)
        spacing(p, before=3, after=0)
        p = doc.add_paragraph(); tabbed(p); ri = p.add_run(org); ri.italic = True
        rl = p.add_run("\t" + loc); rl.italic = True; spacing(p, after=1)
        for b in items:
            bullet(b)

    heading("PROJECTS")
    for title, repo_txt, repo_url, stack, items in PROJECTS:
        p = doc.add_paragraph(); tabbed(p); r = p.add_run(title); r.bold = True
        if repo_url:
            p.add_run("\t"); _add_hyperlink(p, repo_url, "GitHub", size=8.5)
        spacing(p, before=3, after=0)
        p = doc.add_paragraph(); ri = p.add_run(stack); ri.italic = True; ri.font.size = Pt(9); spacing(p, after=1)
        for b in items:
            bullet(b)

    heading("EDUCATION")
    p = doc.add_paragraph(); tabbed(p); r = p.add_run(EDUCATION["title"]); r.bold = True
    p.add_run("\t" + EDUCATION["date"]); spacing(p, before=3, after=0)
    p = doc.add_paragraph(); tabbed(p); ri = p.add_run(EDUCATION["org"]); ri.italic = True
    rr = p.add_run("\t" + EDUCATION["right"]); rr.italic = True; spacing(p, after=1)
    p = doc.add_paragraph(); p.add_run(EDUCATION["extra"]); spacing(p, after=1)

    heading("LEADERSHIP & ACTIVITIES")
    for b in LEADERSHIP:
        bullet(b)

    heading("CERTIFICATIONS")
    p = doc.add_paragraph(); p.add_run(CERTS); spacing(p, after=0)

    doc.save(path)
    return path


if __name__ == "__main__":
    pdf = build_pdf()
    docx_path = build_docx()
    print("Wrote:", os.path.relpath(pdf, ROOT))
    print("Wrote:", os.path.relpath(docx_path, ROOT))
    try:
        import fitz
        d = fitz.open(pdf)
        print(f"PDF pages: {d.page_count}")
    except Exception as e:
        print("page count check skipped:", e)
